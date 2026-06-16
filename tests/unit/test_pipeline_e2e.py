"""End-to-end index/status pipeline over real Docling conversions.

PDF conversion needs HuggingFace model downloads and is excluded here
(user-verify item in the M2 brief); md/txt/docx exercise the real path.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from corpusqa.config import load_config
from corpusqa.ingest.discovery import DriftKind
from corpusqa.ingest.pipeline import index_paths, run_index, run_status

EXAMPLE = Path(__file__).resolve().parents[2] / "corpusqa.example.yaml"


@pytest.fixture(scope="module")
def corpus(tmp_path_factory: pytest.TempPathFactory) -> Path:
    root = tmp_path_factory.mktemp("corpus")
    (root / "notes.md").write_text(
        "# Brainstorm\n\n## Ideas\n" + "An idea about gardens. " * 30,
        encoding="utf-8",
    )
    (root / "plain.txt").write_text("Plain facts about rivers.", encoding="utf-8")
    try:
        import docx as pydocx
    except ImportError:
        pytest.skip("python-docx unavailable for fixture creation")
    d = pydocx.Document()
    d.add_heading("Report", 0)
    d.add_paragraph("Engines and pistons, a study. " * 20)
    d.save(root / "report.docx")
    return root


@pytest.mark.slow
def test_index_then_status_then_drift(corpus: Path) -> None:
    config = load_config(EXAMPLE)

    report = run_index(corpus, config)
    assert report.parsed == 3 and not report.flagged
    assert report.chunks_written >= 3

    _, _, cache = index_paths(corpus)
    md_files = list(cache.glob("*.md"))
    assert len(md_files) == 3
    assert all((cache / f"{p.stem}.chunks.jsonl").exists() for p in md_files)

    status = run_status(corpus, config)
    assert status.file_count == 3 and not status.drift

    # second run is a no-op (incrementality)
    again = run_index(corpus, config)
    assert again.parsed == 0 and again.unchanged == 3

    # drift: modify one, delete one, add one
    (corpus / "plain.txt").write_text("Plain facts, amended.", encoding="utf-8")
    (corpus / "notes.md").unlink()
    (corpus / "fresh.md").write_text("# New\nfresh content", encoding="utf-8")
    drift = {(d.rel_path, d.drift) for d in run_status(corpus, config).drift}
    assert ("plain.txt", DriftKind.MODIFIED) in drift
    assert ("notes.md", DriftKind.DELETED) in drift
    assert ("fresh.md", DriftKind.NEW) in drift

    final = run_index(corpus, config)
    assert final.parsed == 2 and final.deleted == 1
    assert run_status(corpus, config).file_count == 3
